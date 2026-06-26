"""Utility API — file upload/extract, export, RAG document management."""

import io
import logging
from fastapi import APIRouter, File, Form, UploadFile, HTTPException, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document

from app.db import get_db, User, Document
from app.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(tags=["utils"])

ALLOWED_EXTENSIONS = {".docx", ".pdf", ".txt"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB

# Magic bytes for file type verification
DOCX_MAGIC = b"PK\x03\x04"
PDF_MAGIC = b"%PDF"


def _verify_file_type(content: bytes, ext: str):
    """Verify file content matches its claimed extension via magic bytes."""
    if ext == ".docx":
        # .docx is a ZIP archive with PK header
        if not content.startswith(DOCX_MAGIC):
            raise HTTPException(status_code=400, detail="文件格式与扩展名不匹配，请确认是有效的 .docx 文件")
    elif ext == ".pdf":
        if not content.startswith(PDF_MAGIC):
            raise HTTPException(status_code=400, detail="文件格式与扩展名不匹配，请确认是有效的 .pdf 文件")


def extract_text_from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def extract_text_from_pdf(content: bytes) -> str:
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    try:
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    finally:
        doc.close()


def extract_text_from_txt(content: bytes) -> str:
    # Try common Chinese encodings in order
    for encoding in ("utf-8", "gbk", "gb2312", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


@router.post("/upload/extract-text")
async def extract_text(file: UploadFile = File(...)):
    if not file.filename or not file.filename.strip():
        raise HTTPException(status_code=400, detail="文件名不能为空")

    ext = file.filename.rsplit(".", 1)[-1].lower().strip()
    ext = f".{ext}"
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {ext}，仅支持 .docx / .pdf / .txt")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="文件内容为空")
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"文件过大，最大支持 {MAX_FILE_SIZE // 1024 // 1024}MB")

    _verify_file_type(content, ext)

    try:
        if ext == ".docx":
            text = extract_text_from_docx(content)
        elif ext == ".pdf":
            text = extract_text_from_pdf(content)
        else:
            text = extract_text_from_txt(content)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Extract error: {e}")
        raise HTTPException(status_code=500, detail="文件解析失败，请确认文件未损坏且格式正确")

    if not text.strip():
        raise HTTPException(status_code=400, detail="未能从文件中提取到文字内容")

    return {"filename": file.filename, "text": text, "length": len(text)}


@router.post("/export/docx")
async def export_docx(text: str = Form(..., min_length=1)):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.size = 157000  # 12pt in EMU

    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = 157000

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=rewritten.docx"},
    )


# ── RAG document management (auth required) ──

@router.post("/rag/upload")
async def rag_upload(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if not file.filename or not file.filename.strip():
        raise HTTPException(status_code=400, detail="文件名不能为空")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("pdf", "docx", "txt"):
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: .{ext}，仅支持 .pdf / .docx / .txt")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="文件内容为空")
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件过大，最大支持 20MB")

    try:
        from app.rag.manager import get_rag_manager
        mgr = get_rag_manager()
        result = mgr.index_document(content, file.filename)

        # Save metadata to DB
        doc_record = Document(
            user_id=user.id,
            filename=file.filename,
            chunk_count=result["chunk_count"],
            char_count=result["char_count"],
        )
        db.add(doc_record)
        await db.commit()

        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"RAG upload error: {e}")
        raise HTTPException(status_code=500, detail="文献索引失败")


@router.get("/rag/documents")
async def rag_list(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy import select as slc
    result = await db.execute(slc(Document).where(Document.user_id == user.id))
    records = result.scalars().all()
    return [
        {
            "id": r.id,
            "filename": r.filename,
            "chunk_count": r.chunk_count,
            "char_count": r.char_count,
            "created_at": r.created_at.isoformat() if r.created_at else "",
        }
        for r in records
    ]


@router.delete("/rag/documents/{doc_id}")
async def rag_delete(
    doc_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy import select as slc
    result = await db.execute(slc(Document).where(Document.id == doc_id, Document.user_id == user.id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文献记录不存在")

    from app.rag.manager import get_rag_manager
    mgr = get_rag_manager()
    # Remove vector data (we only have by document_id hash, skip for now)

    await db.delete(doc)
    await db.commit()
    return {"ok": True}
